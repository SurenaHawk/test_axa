from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from .models import Projet, PlanImage
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.conf import settings
import os
from django.core.files.base import ContentFile
from datetime import datetime, timedelta
from bs4 import BeautifulSoup
import tempfile

# Create your views here.

def home(request):
    projets = Projet.objects.all()
    return render(request, 'contrat/home.html', {'projets': projets})

@csrf_exempt    
def devis(request):
    if request.method == 'POST':
        opportunity_number = request.POST.get("opportunity_number")
        client_name = request.POST.get("client_name")
        siret_number = request.POST.get("siret_number")
        siren_number = request.POST.get("siren_number")
        affaire = request.POST.get("affaire")
        reference = request.POST.get("reference")
        intermediaire = request.POST.get("intermediaire")
        short_description = request.POST.get("short_description")
        image = request.FILES.get("image_description")
        coassurance = request.POST.get("coassurance")
        operation_adresse = request.POST.get("operation_adresse")
        plan_operation = request.FILES.getlist("plan_operation")
        operation_description = request.POST.get("operation_description")
        price_operation = request.POST.get('tarif')

        if (
            opportunity_number is not None and client_name is not None and siret_number is not None and siren_number is not None and affaire is not None 
            and reference is not None and intermediaire is not None and short_description is not None and image is not None and coassurance is not None 
            and operation_adresse is not None and plan_operation is not None and operation_description is not None and price_operation is not None
        ):
            coassurance_field = False
            if coassurance == "True":
                coassurance_field = True
            projet = Projet.objects.create(
                opportunity_number = opportunity_number,
                client_name = client_name,
                siret_number = siret_number,
                siren_number = siren_number,
                affaire = affaire,
                reference = reference,
                intermediaire = intermediaire,
                short_description = short_description,
                image = image,
                coassurance = coassurance_field,
                operation_adresse = operation_adresse,
                operation_description = operation_description,
                price_operation = price_operation
            )

            projet.save()

            for image in plan_operation:
                projet_image = PlanImage(projet=projet, image_plan_operation=image)
                projet_image.save()

            data = [
                opportunity_number, client_name, siret_number, siren_number, affaire, reference, intermediaire, short_description, 
                image, coassurance, operation_adresse, operation_description, price_operation
            ]

            current_datetime = datetime.now() + timedelta(hours=2)
            pdf_content = generate_pdf(data)
            docx_content = generate_docx(data)
            projet.pdf_file.save('Projet_de_contrat_'+opportunity_number+'_Date_'+current_datetime.strftime('%Y-%m-%d_%H-%M-%S')+'.pdf', ContentFile(pdf_content))
            projet.docx_file.save('Projet_de_contrat_'+opportunity_number+'_Date_'+current_datetime.strftime('%Y-%m-%d_%H-%M-%S')+'.docx', ContentFile(docx_content))

            message = {'message': 'Le devis a bien été créé. Veuillez cliquer sur le menu "Accueil" ou sur le logo pour être redirigé vers la page d\'accueil', 
                       'status':200}
            return JsonResponse(message, safe=False)

    else:
        return render(request, 'contrat/devis.html')
    return render(request, 'contrat/devis.html')

def generate_pdf(data):
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    
    left_margin = 50
    top_margin = 600

    p.drawImage("https://upload.wikimedia.org/wikipedia/commons/thumb/9/94/AXA_Logo.svg/1200px-AXA_Logo.svg.png", left_margin, top_margin, 100,80)
    
    p.setFont("Helvetica-Bold", 14)
    devis_width = p.stringWidth("Devis", "Helvetica-Bold")
    p.drawString(letter[0] - left_margin - devis_width, 750, "Devis")

    p.setFont("Helvetica-Bold", 10)
    p.drawString(50, 620 - 40 - 20, "Numéro opportunité : {}".format(data[0]))
    p.drawString(50, 620 - 40 - 40, "Référence : {}".format(data[5]))
    p.drawString(50, 620 - 40 - 60, "Nom client : {}".format(data[1]))
    p.drawString(50, 620 - 40 - 80, "Intermédiaire : {}".format(data[6]))
    p.drawString(50, 620 - 40 - 100, "Description succinte : {}".format(data[7]))
    p.drawString(50, 620 - 40 - 120, "Image description : ")
    image_description = b''.join(chunk for chunk in data[8].chunks())
    with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp_file:
        tmp_file.write(image_description)
        tmp_file_path = tmp_file.name
    p.drawImage(tmp_file_path, 50, 620-40-180, width=2*inch, height=1*inch)

    co_assurance = "Oui"
    if data[9] == 'false':
        co_assurance = "Non"
    p.drawString(50, 620 - 40 - 200, "Co assurance : {}".format(co_assurance))
    p.drawString(50, 620 - 40 - 220, "Adresse opération : {}".format(data[10]))

    p.drawString(50, 620 - 40 - 240, "Description opération : ")
    p.drawString(50, 620 - 40 - 260, BeautifulSoup(data[11], "html.parser").get_text(separator=' '))
    p.drawString(50, 620 - 40 - 280, "Tarif : {}€".format(data[12]))

    p.showPage()
    p.save()
    pdf_content = buffer.getvalue()
    buffer.close()
    return pdf_content

def generate_docx(data):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.top_margin = Inches(1)
    
    devis = doc.add_paragraph()
    devis.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    devis_text = devis.add_run("Devis")
    devis_text.bold = True

    logo_path = os.path.join(settings.BASE_DIR, 'contrat/static', 'img', 'axa_logo.png')
    doc.add_picture(logo_path, width=Inches(1.0))  
    
    client_name = doc.add_paragraph()
    client_text = client_name.add_run("Nom du Client : ")
    client_text.bold = True
    client_name.add_run(data[1])
    client_name.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    opportunity = doc.add_paragraph()
    opportunity_text = opportunity.add_run("Numéro opportunité : ")
    opportunity_text.bold = True
    opportunity.add_run(data[0])

    reference = doc.add_paragraph()
    reference_text = reference.add_run("Référence : ")
    reference_text.bold = True
    reference.add_run(data[5])

    intermediare = doc.add_paragraph()
    intermediare_text = intermediare.add_run("Intermédiaire : ")
    intermediare_text.bold = True
    intermediare.add_run(data[6])

    description = doc.add_paragraph()
    description_text = description.add_run("Description succinte : ")
    description_text.bold = True
    description.add_run(data[7])

    image_description = doc.add_paragraph()
    image_description_text = image_description.add_run("Image description : ")
    image_description_text.bold = True
    image_data = b''.join(chunk for chunk in data[8].chunks())
    image_description = BytesIO(image_data)
    doc.add_picture(image_description, width=Inches(2.0))

    coassurance = doc.add_paragraph()
    coassurance_text = coassurance.add_run("Co assurance : ")
    coassurance_text.bold = True
    if data[9] == "true":
        coassurance.add_run("Oui")
    else:
        coassurance.add_run("Non")

    adresse_opeation = doc.add_paragraph()
    adresse_opeation_text = adresse_opeation.add_run("Adresse opération : ")
    adresse_opeation_text.bold = True
    adresse_opeation.add_run(data[10])

    description_operation = doc.add_paragraph()
    description_operation_text = description_operation.add_run("Description opération : ")
    description_operation_text.bold = True
    description_operation = doc.add_paragraph()
    description_operation_text_clean = BeautifulSoup(data[11], "html.parser").get_text()
    description_operation.add_run(description_operation_text_clean)

    tarif = doc.add_paragraph()
    tarif_text = tarif.add_run("Tarif : €")
    tarif_text.bold = True
    tarif.add_run(data[12])

    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_content = docx_buffer.getvalue()
    docx_buffer.close()
    return docx_content

def download_pdf(request):
    opportunity_number = request.GET.get('opportunity_number')
    projets = Projet.objects.filter(opportunity_number=opportunity_number)
    current_datetime = datetime.now() + timedelta(hours=2)
    for projet in projets:
        pdf_document = projet.pdf_file
        response = HttpResponse(pdf_document, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Projet_de_contrat_'+opportunity_number+'_Date_'+current_datetime.strftime('%Y-%m-%d_%H-%M-%S')+'.pdf'
    return response

def download_docx(request):
    opportunity_number = request.GET.get('opportunity_number')
    projets = Projet.objects.filter(opportunity_number=opportunity_number)
    current_datetime = datetime.now() + timedelta(hours=2)
    for projet in projets:
        pdf_document = projet.docx_file
        response = HttpResponse(pdf_document, content_type='application/docx')
        response['Content-Disposition'] = f'attachment; filename="Projet_de_contrat_'+opportunity_number+'_Date_'+current_datetime.strftime('%Y-%m-%d_%H-%M-%S')+'.docx'
    return response
